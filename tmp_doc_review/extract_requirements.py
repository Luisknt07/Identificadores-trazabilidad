import json
import sys
from pathlib import Path

from docx import Document
from docx.table import Table

sys.stdout.reconfigure(encoding="utf-8")
from docx.text.paragraph import Paragraph


def iter_blocks(document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


source = Path(sys.argv[1])
document = Document(source)
blocks = []
for block in iter_blocks(document):
    if isinstance(block, Paragraph):
        text = block.text.strip()
        if text:
            blocks.append({"type": "paragraph", "style": block.style.name if block.style else "", "text": text})
    else:
        rows = []
        for row in block.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        blocks.append({"type": "table", "rows": rows})

payload = {
    "paragraph_count": len(document.paragraphs),
    "table_count": len(document.tables),
    "inline_shape_count": len(document.inline_shapes),
    "section_count": len(document.sections),
    "blocks": blocks,
}
if len(sys.argv) > 2:
    lines = []
    for block in blocks:
        if block["type"] == "paragraph":
            style = block["style"]
            text = block["text"]
            if style.startswith("Heading"):
                level = style.rsplit(" ", 1)[-1]
                prefix = "#" * int(level) if level.isdigit() else "#"
                lines.extend([f"{prefix} {text}", ""])
            elif style == "Source Code":
                lines.extend(["```text", text, "```", ""])
            else:
                lines.extend([text, ""])
        else:
            for row in block["rows"]:
                lines.append(" | ".join(row))
            lines.append("")
    Path(sys.argv[2]).write_text("\n".join(lines), encoding="utf-8")
else:
    print(json.dumps(payload, ensure_ascii=False, indent=2))
